"""
Document extraction service for PDF/Office file content.

Provides methods to extract text from documents for search indexing.
"""
import hashlib
import os
from typing import Optional
from pathlib import Path

from django.contrib.contenttypes.models import ContentType


class DocumentService:
    """
    Service for extracting text from documents (PDFs, Word, Excel).

    Uses PyPDF2 for PDFs, python-docx for Word documents,
    and openpyxl for Excel files.
    """

    def __init__(self):
        self._pypdf2_available = None
        self._docx_available = None
        self._openpyxl_available = None

    @property
    def pypdf2_available(self) -> bool:
        """Check if PyPDF2 is available."""
        if self._pypdf2_available is None:
            try:
                import PyPDF2
                self._pypdf2_available = True
            except ImportError:
                self._pypdf2_available = False
        return self._pypdf2_available

    @property
    def docx_available(self) -> bool:
        """Check if python-docx is available."""
        if self._docx_available is None:
            try:
                import docx
                self._docx_available = True
            except ImportError:
                self._docx_available = False
        return self._docx_available

    @property
    def openpyxl_available(self) -> bool:
        """Check if openpyxl is available."""
        if self._openpyxl_available is None:
            try:
                import openpyxl
                self._openpyxl_available = True
            except ImportError:
                self._openpyxl_available = False
        return self._openpyxl_available

    def get_file_checksum(self, file_path: str) -> str:
        """Calculate MD5 checksum of a file."""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except (IOError, OSError):
            return ""

    def detect_file_type(self, file_path: str) -> str:
        """Detect the file type from extension."""
        ext = Path(file_path).suffix.lower()
        mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'docx',
            '.xlsx': 'xlsx',
            '.xls': 'xlsx',
            '.txt': 'txt',
        }
        return mapping.get(ext, 'other')

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a document file.

        Automatically detects file type and uses appropriate extractor.
        Returns empty string if extraction fails.
        """
        if not os.path.exists(file_path):
            return ""

        file_type = self.detect_file_type(file_path)

        if file_type == 'pdf':
            return self._extract_pdf(file_path)
        elif file_type == 'docx':
            return self._extract_docx(file_path)
        elif file_type == 'xlsx':
            return self._extract_xlsx(file_path)
        elif file_type == 'txt':
            return self._extract_txt(file_path)

        return ""

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        if not self.pypdf2_available:
            return ""

        try:
            import PyPDF2

            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return ' '.join(text_parts)
        except Exception:
            return ""

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from a Word document."""
        if not self.docx_available:
            return ""

        try:
            import docx

            doc = docx.Document(file_path)
            text_parts = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            return ' '.join(text_parts)
        except Exception:
            return ""

    def _extract_xlsx(self, file_path: str) -> str:
        """Extract text from an Excel spreadsheet."""
        if not self.openpyxl_available:
            return ""

        try:
            import openpyxl

            wb = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            text_parts.append(str(cell.value))

            wb.close()
            return ' '.join(text_parts)
        except Exception:
            return ""

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from a plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""

    def index_digital_asset(self, asset) -> Optional['SearchIndex']:
        """
        Extract and index text from a digital asset.

        Creates or updates a SearchIndex record for the asset.
        """
        from ..models import SearchIndex, SearchSettings

        settings = SearchSettings.get_settings()
        if not settings.index_documents:
            return None

        # Get file path from asset
        if not hasattr(asset, 'file') or not asset.file:
            return None

        file_path = asset.file.path

        # Check checksum to see if reindex is needed
        new_checksum = self.get_file_checksum(file_path)
        file_type = self.detect_file_type(file_path)

        # Get content type for DigitalAsset
        content_type = ContentType.objects.get_for_model(asset)

        # Get or create SearchIndex
        search_index, created = SearchIndex.objects.get_or_create(
            content_type=content_type,
            object_id=asset.pk,
            defaults={
                'file_type': file_type,
                'checksum': new_checksum,
            }
        )

        # Check if needs reindex
        if not created and search_index.checksum == new_checksum:
            return search_index

        # Extract text
        extracted_text = self.extract_text(file_path)

        # Update SearchIndex
        search_index.extracted_text = extracted_text
        search_index.file_type = file_type
        search_index.checksum = new_checksum
        search_index.save()

        return search_index

    def needs_reindex(self, asset) -> bool:
        """
        Check if a digital asset needs re-indexing.

        Returns True if the file has changed since last indexing.
        """
        from ..models import SearchIndex

        if not hasattr(asset, 'file') or not asset.file:
            return False

        try:
            content_type = ContentType.objects.get_for_model(asset)
            search_index = SearchIndex.objects.get(
                content_type=content_type,
                object_id=asset.pk
            )

            current_checksum = self.get_file_checksum(asset.file.path)
            return search_index.checksum != current_checksum
        except SearchIndex.DoesNotExist:
            return True

    def bulk_index_assets(self, assets, batch_size: int = 50) -> dict:
        """
        Bulk index multiple digital assets.

        Returns dict with counts of indexed, skipped, and failed.
        """
        results = {
            'indexed': 0,
            'skipped': 0,
            'failed': 0,
        }

        for asset in assets:
            try:
                if not self.needs_reindex(asset):
                    results['skipped'] += 1
                    continue

                index = self.index_digital_asset(asset)
                if index:
                    results['indexed'] += 1
                else:
                    results['skipped'] += 1
            except Exception:
                results['failed'] += 1

        return results

    def search_document_content(self, query: str, limit: int = 10) -> list:
        """
        Search through indexed document content.

        Returns list of matching SearchIndex records.
        """
        from ..models import SearchIndex

        return list(
            SearchIndex.objects.filter(
                extracted_text__icontains=query
            )[:limit]
        )

    def get_document_snippet(self, text: str, query: str,
                             snippet_length: int = 200) -> str:
        """
        Get a snippet of text around the query match.

        Returns snippet with the query highlighted.
        """
        query_lower = query.lower()
        text_lower = text.lower()

        # Find query position
        pos = text_lower.find(query_lower)
        if pos == -1:
            return text[:snippet_length] + '...' if len(text) > snippet_length else text

        # Calculate snippet boundaries
        start = max(0, pos - snippet_length // 2)
        end = min(len(text), pos + len(query) + snippet_length // 2)

        snippet = text[start:end]

        # Add ellipsis if truncated
        if start > 0:
            snippet = '...' + snippet
        if end < len(text):
            snippet = snippet + '...'

        return snippet
